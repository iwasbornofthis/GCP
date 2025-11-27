# -*- coding: utf-8 -*-
from docx import Document

def add_list(doc, items, style):
    for item in items:
        try:
            doc.add_paragraph(item, style=style)
        except Exception as e:
            raise RuntimeError(f'Failed item: {item!r} style={style}: {e}')

doc = Document()
doc.add_paragraph('test')
data_flow = ['① test', '② next']
add_list(doc, data_flow, 'List Number')
