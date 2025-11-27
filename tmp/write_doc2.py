# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_list(doc, items, style):
    for item in items:
        doc.add_paragraph(item, style=style)

doc = Document()

doc.add_heading('프로젝트 이름', level=0)
doc.add_paragraph('AI 기반 식품 분석 · 건강 코치 서비스 "FoodScan"', style='Subtitle')
doc.add_paragraph('요구사항 명세서 및 사업계획서', style='Subtitle')

doc.add_paragraph('프로젝트 버전: 1.0')
doc.add_paragraph('작성일: 2025년 10월 30일')

doc.add_heading('목차', level=1)
toc = doc.add_paragraph()
toc.add_run('1. 프로젝트 개요 / 2. 요구사항 명세 / 3. 시스템 설계 / 4. 기능 명세 / 5. 사용자 시나리오 / 6. 개발 일정 / 7. 기대 효과 / 8. 위험 관리 / 9. 기술 스택 / 10. 참고 문헌')
toc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Section 1
doc.add_heading('1. 프로젝트 개요', level=1)
doc.add_heading('1.1 프로젝트 배경', level=2)
doc.add_paragraph('초개인화 식습관 관리 수요가 증가했지만, 사용자는 음식 기록과 영양 계산을 직접 해야 하는 불편을 겪고 있습니다. 스마트폰 카메라와 대규모 언어모델을 활용하면 사진 한 장으로 정확한 영양 분석과 권장 섭취 가이드를 제공할 수 있습니다.')

doc.add_heading('1.2 프로젝트 목적', level=2)
add_list(doc, [
    '사진 기반 음식 인식과 DB 매칭으로 즉시 영양 정보를 보여주는 웹 서비스를 구축한다.',
    'OpenAI GPT-4o와 Food Matcher 임베딩을 결합해 한국형 식품 데이터에서도 높은 정확도를 확보한다.',
    '사용자 프로필·섭취 이력을 이용해 맞춤 리포트와 건강 목표 달성 힌트를 제공한다.'
], 'List Bullet')

doc.add_heading('1.3 프로젝트 범위', level=2)
add_list(doc, [
    'React 기반 모바일 우선 UI와 QR/카메라 모드 지원',
    'Laravel 12 REST API, 인증/토큰, 이미지 업로드, OpenAI·FastAPI 연동',
    'SentenceTransformers 임베딩을 사용하는 Food Matcher FastAPI',
    'MySQL 영양 데이터 정제, product_weight 환산, common_name 평균 보정'
], 'List Bullet')

# Section 2
doc.add_heading('2. 요구사항 명세', level=1)
doc.add_heading('2.1 기능적 요구사항 (Functional Requirements)', level=2)
add_list(doc, [
    'F1. 사용자는 이메일/비밀번호로 회원가입·로그인 후 토큰을 발급받는다.',
    'F2. 사진 또는 갤러리 이미지를 업로드하면 OpenAI를 통해 음식명을 추출한다.',
    'F3. Food Matcher는 음식 설명을 임베딩 해 DB 내 최적의 food_code를 반환한다.',
    'F4. API는 영양소, 알레르기, 추천 문구를 JSON으로 응답하고 React UI가 카드를 표시한다.',
    'F5. 사용자 프로필·섭취 로그·알림 이력을 DB에 저장해 추후 리포트에 활용한다.'
], 'List Number')

doc.add_heading('2.2 비기능적 요구사항 (Non-Functional Requirements)', level=2)
add_list(doc, [
    'N1. API 응답 시간은 5초 이내, OpenAI 지연 시 재시도 및 타임아웃 처리',
    'N2. 사용자 사진·토큰은 HTTPS와 JWT로 보호하고, 민감 정보는 암호화한다.',
    'N3. FastAPI와 Laravel 로그를 통합 수집해 오류를 추적한다.',
    'N4. Docker 기반 로컬 개발 환경을 제공하고, staging/prod 분리 구성',
    'N5. 한국어 기본 UI와 영어 fallback을 제공한다.'
], 'List Number')

# Section3
doc.add_heading('3. 시스템 설계', level=1)
doc.add_heading('3.1 시스템 아키텍처', level=2)
doc.add_paragraph('FoodScan은 React 프런트엔드, Laravel API, FastAPI Food Matcher, MySQL DB로 구성된 3-Tier 구조입니다. React는 Vite 개발 서버와 ngrok 도메인을 사용해 HTTPS/HTTP 모드를 전환하고, Laravel은 OpenAI 및 매칭 서비스와 통신하는 BFF 역할을 수행합니다.')

doc.add_heading('3.2 데이터 흐름도', level=2)
add_list(doc, [
    '사용자가 사진을 촬영하거나 업로드해 React에서 /api/foods/analyze 로 전송한다.',
    'Laravel이 이미지 유효성 검증 후 OpenAI GPT-4o 멀티모달 API를 호출한다.',
    '생성된 음식 설명을 FastAPI Food Matcher로 전달해 food_code와 score를 획득한다.',
    'MySQL foods 테이블에서 영양 정보를 조회하고 정제된 값을 조합한다.',
    'JSON 응답을 React가 수신해 결과 카드, 오류 안내, 추천 문구를 표시한다.'
], 'List Number')

doc.add_heading('3.3 데이터베이스 스키마', level=2)
doc.add_paragraph('주요 테이블은 USERS, MEAL_RECORDS, FOODS, FOOD_EMBEDDINGS입니다. FOODS 테이블은 common_name 평균과 product_weight 환산을 통해 결측값을 최소화했고, FOOD_EMBEDDINGS는 SentenceTransformers 임베딩을 저장해 매칭 속도를 높였습니다.')

# Section4
doc.add_heading('4. 기능 명세', level=1)
doc.add_heading('4.1 사용자 프로필 관리', level=2)
doc.add_paragraph('회원가입 시 나이, 성별, 키, 체중, 목표 칼로리를 입력받고, Laravel Sanctum 토큰으로 보호합니다. 프로필은 React 설정 화면에서 수정하고 DB에 즉시 반영됩니다.')

doc.add_heading('4.2 AI 식단 추천 엔진', level=2)
doc.add_paragraph('OpenAI GPT-4o의 멀티모달 기능으로 음식명과 문장형 설명을 생성하고, Food Matcher가 임베딩 유사도로 국내 식품 데이터를 찾습니다. 점수가 낮을 경우 대안 후보를 최대 5개까지 반환해 UI에서 선택할 수 있습니다.')

doc.add_heading('4.3 식단 기록 및 분석', level=2)
doc.add_paragraph('MEAL_RECORDS 테이블에 섭취 시간, 이미지 경로, food_code, 영양소 합계를 저장합니다. React는 일별·주별 통계를 표시하고, Laravel은 열량 초과 시 경고 메시지를 제공합니다.')

doc.add_heading('4.4 알림 및 피드백 시스템', level=2)
doc.add_paragraph('분석 결과는 푸시 또는 이메일 템플릿으로 확장 가능하도록 설계했습니다. 목표 대비 80% 미만·120% 초과 시 알림을 전송하고, 향후 기기 연동을 위한 Hook을 정의했습니다.')

# Section5
doc.add_heading('5. 사용자 시나리오', level=1)
doc.add_heading('5.1 신규 사용자 온보딩 시나리오', level=2)
add_list(doc, [
    '사용자가 QR 코드로 FoodScan에 접속해 회원가입한다.',
    '프로필 정보를 입력하고 카메라 권한을 허용한다.',
    '튜토리얼에서 사진 분석 절차를 체험한다.',
    '첫 분석 결과와 피드백을 수신한다.'
], 'List Number')

doc.add_heading('5.2 식단 추천 시나리오', level=2)
add_list(doc, [
    '사용자가 점심 사진을 업로드한다.',
    'OpenAI와 Food Matcher가 유사 음식 3건과 영양소를 반환한다.',
    'React가 칼로리·탄수화물 과다 여부를 표시하고 대체 식단을 제안한다.',
    '사용자가 추천 레시피 링크와 저장 버튼을 선택한다.'
], 'List Number')

doc.add_heading('5.3 식단 기록 및 피드백 시나리오', level=2)
add_list(doc, [
    '사용자가 하루 섭취 기록을 확인하고 메모를 추가한다.',
    'Laravel이 주간 리포트를 생성해 이메일 또는 알림으로 전달한다.',
    'FastAPI 점수와 DB 통계를 합쳐 eGL 예측 준비 데이터를 축적한다.'
], 'List Number')

# Section6
doc.add_heading('6. 개발 일정', level=1)
doc.add_heading('6.1 간트 차트', level=2)
doc.add_paragraph('10월: 요구사항 정의 및 DB 정제 / 11월: 프런트·백엔드 핵심 기능 개발 / 12월: AI 파이프라인 안정화 및 테스트, 운영 가이드 작성.')

doc.add_heading('6.2 주차별 상세 계획', level=2)
add_list(doc, [
    'W1~W2: 데이터 정제, product_weight 환산, common_name 평균 적용',
    'W3~W4: React UI, 카메라·갤러리 모드 완성, OpenAI 연동',
    'W5~W6: Food Matcher FastAPI, embed_foods 파이프라인 구축',
    'W7~W8: 통합 테스트, 보안 점검, 운영 문서 업데이트'
], 'List Bullet')

# Section7
doc.add_heading('7. 기대 효과', level=1)
doc.add_heading('7.1 정량적 목표', level=2)
add_list(doc, [
    '분석 성공률 90% 이상, 평균 응답 시간 4초 이하',
    '주간 활성 사용자 1,000명, 재방문율 60% 달성',
    'DB 결측률 5% 이하 유지'
], 'List Number')

doc.add_heading('7.2 정성적 효과', level=2)
doc.add_heading('사용자 측면', level=3)
doc.add_paragraph('사진 한 장으로 영양소를 확인하고 목표 대비 현황을 즉시 파악할 수 있어 자기주도적 식습관 개선을 돕습니다.')

doc.add_heading('서비스 측면', level=3)
doc.add_paragraph('OpenAI와 Food Matcher 기반 모듈형 구조로 다른 서비스와 연동해 부가가치를 창출할 수 있습니다.')

doc.add_heading('시장 측면', level=3)
doc.add_paragraph('국내 식품 DB를 활용한 AI 영양 코치 서비스로 헬스케어 시장에서 차별화 포지셔닝이 가능합니다.')

doc.add_heading('7.3 기대 효과 시각화', level=2)
doc.add_paragraph('대시보드에는 일별 분석 건수, 평균 응답시간, 권장 섭취 달성률을 시계열 그래프로 표시하고, 사용자 세그먼트별 만족도를 히트맵으로 표현합니다.')

# Section8
doc.add_heading('8. 위험 관리', level=1)
doc.add_heading('8.1 잠재적 위험 요소 및 대응 방안', level=2)
add_list(doc, [
    'OpenAI API 한도 초과 -> 요청 큐 관리, 이미지 크기 제한, 대체 모델 준비',
    '식품 DB 오류 -> 정기 크롤링, 수동 검수, 평균 보정 로그 유지',
    '개인정보 유출 -> HTTPS 강제, JWT 만료, 최소 데이터 수집 원칙'
], 'List Bullet')

doc.add_heading('8.2 품질 보증 계획(이 부분은 참조만 하세요)', level=2)
doc.add_paragraph('React 단위 테스트, Laravel 기능 테스트, FastAPI pytest를 모두 CI에 포함합니다. OWASP ZAP으로 월 1회 취약점 점검을 실시합니다.')

# Section9
doc.add_heading('9. 기술 스택', level=1)
doc.add_heading('9.1 개발 환경', level=2)
doc.add_paragraph('Node.js 18, npm 10, PHP 8.2, Composer 2, Python 3.11, MySQL 8, Redis 6, Windows 11 개발 머신(XAMPP).')

doc.add_heading('9.2 기술 스택 아키텍처', level=2)
doc.add_paragraph('프런트: React, Vite, Tailwind, jsQR / 백엔드: Laravel, Sanctum, Guzzle / AI: OpenAI GPT-4o, SentenceTransformers, FastAPI / 데이터: MySQL, Redis 캐시, S3(이미지 저장 예정).')

# Section10
doc.add_heading('10. 참고 문헌', level=1)
add_list(doc, [
    'Laravel 12 공식 문서 (https://laravel.com/docs/12.x)',
    'OpenAI API Reference (https://platform.openai.com/docs)',
    'SentenceTransformers Documentation (https://www.sbert.net)',
    '식품안전나라 영양성분 데이터베이스 (https://www.foodsafetykorea.go.kr)',
    'OWASP API Security Top 10, 2023'
], 'List Bullet')


doc.save(r'C:\AI\de_new.docx')
