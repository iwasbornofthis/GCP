# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

doc.add_heading('FoodScan 기능 명세서', level=0)
intro = doc.add_paragraph('이 문서는 React 기반 FoodScan 웹 애플리케이션과 Laravel API, Food Matcher 임베딩 서비스, 데이터베이스 정제 흐름을 포함한 전체 기능을 개괄합니다. 향후 발표 자료(PPT) 작성 시 참고하도록 실제 구현 내용을 중심으로 정리했습니다.')
intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_heading('1. 프로젝트 개요', level=1)
doc.add_paragraph('프로젝트명: FoodScan – AI 기반 식품 영양 분석 및 추천 서비스')
doc.add_paragraph('기간: 2025년 10월 ~ 2025년 12월 / 버전: v1.0')

doc.add_heading('1.1 목표', level=2)
doc.add_paragraph('· 사진 한 장으로 음식명을 추출하고 영양소·권장 정보를 즉시 보여주는 모바일 최적화 웹을 제공한다.')
doc.add_paragraph('· 사용자가 누적한 섭취 데이터를 내부 DB와 연계하여 개인화된 분석·추천 기능을 확장할 수 있는 기반을 마련한다.')
doc.add_paragraph('· React, Laravel, FastAPI, OpenAI를 연계한 End-to-End AI 파이프라인을 구축한다.')

doc.add_heading('1.2 범위', level=2)
doc.add_paragraph('· React 프런트엔드: 카메라/갤러리 업로드, QR 스캔, 스캔 결과 UI, 오류 안내')
doc.add_paragraph('· Laravel API: 인증, 이미지 업로드, OpenAI 연동, Food Matcher 연동, 영양 정보 응답')
doc.add_paragraph('· Food Matcher FastAPI: SentenceTransformers 임베딩 기반 음식 유사도 계산')
doc.add_paragraph('· MySQL DB: 음식 영양 DB, 정제·전처리 로직, 사용자 계정/토큰 관리')

doc.add_heading('1.3 주요 이해관계자', level=2)
doc.add_paragraph('· 최종 사용자: 모바일 브라우저로 FoodScan에 접속해 사진을 분석하고 결과를 확인하는 일반 사용자')
doc.add_paragraph('· 운영자: DB 정제를 수행하고 OpenAI/Matcher 키를 관리하며 API 가용성을 모니터링하는 관리자')
doc.add_paragraph('· 개발자: React·Laravel 코드 유지보수, FastAPI 임베딩 서비스 개선, 새 모델 실험')

doc.add_heading('2. 시스템 아키텍처', level=1)
doc.add_paragraph('FoodScan은 3계층 구조로 동작합니다.')
for item in ['프런트엔드(React): PWA 스타일 UI, Vite 기반 로컬 개발, env 주입으로 API 주소 제어',
             '백엔드(Laravel 12): REST API, 인증/권한, 이미지 업로드, OpenAI 통신, Food Matcher 호출',
             'AI 매칭(FastAPI): SentenceTransformers 모델과 MySQL에서 추출한 임베딩으로 음식 코드 매칭']:
    doc.add_paragraph(item, style='List Bullet')
workflow = doc.add_paragraph('전체 요청 흐름: 사진 업로드 → Laravel 검증 → OpenAI GPT-4o 멀티모달 요청 → Food Matcher 유사도 계산 → MySQL 영양 정보 조회 → React UI 출력.')
workflow.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_heading('3. 프런트엔드(React)', level=1)
doc.add_heading('3.1 주요 화면', level=2)
doc.add_paragraph('· 촬영 모드: 카메라 권한 요청, 미리보기, 가로/세로 대응, 분석 버튼')
doc.add_paragraph('· 갤러리 모드: 로컬 파일 첨부, drag & drop 대응, 썸네일 프리뷰')
doc.add_paragraph('· 결과 카드: 영양소 표시(칼로리, 탄수화물, 단백질, 지방, 당, 나트륨), 추천 문구, 오류 메시지')

doc.add_heading('3.2 기술 포인트', level=2)
doc.add_paragraph('· Vite 환경변수(VITE_API_BASE_URL)로 API 주소 주입, 개발 시 /api 프록시 사용')
doc.add_paragraph('· HTTPS/HTTP 전환 플래그(VITE_ENABLE_HTTPS) 및 VITE_ALLOWED_HOSTS로 ngrok 도메인 허용')
doc.add_paragraph('· jsQR를 이용한 QR 모드, requestAnimationFrame으로 스캔 라인 애니메이션 구현')

doc.add_heading('4. 백엔드(Laravel 12)', level=1)
doc.add_heading('4.1 인증 및 사용자', level=2)
doc.add_paragraph('· POST /api/register, POST /api/login, GET /api/me, POST /api/logout 로 토큰 기반 인증 제공')
doc.add_paragraph('· 기본 시드 사용자: test@example.com / password')

doc.add_heading('4.2 음식 분석 API', level=2)
doc.add_paragraph('· POST /api/foods/analyze 는 이미지 파일을 수신하고 OpenAI에 전달해 음식 라벨을 생성한다.')
doc.add_paragraph('· Food Matcher FastAPI에 음식 설명을 전달해 가장 유사한 food_code 를 찾고, MySQL foods 테이블에서 영양 정보를 조회한다.')
doc.add_paragraph('· 응답 JSON에는 foodName, confidence, matchedFood, nutrients(칼로리, 단백질, 지방, 탄수화물, 당, 나트륨 등)이 포함된다.')

doc.add_heading('4.3 OpenAI 연동', level=2)
doc.add_paragraph('· .env 에 OPENAI_API_KEY, OPENAI_MODEL=gpt-4o-mini, OPENAI_TIMEOUT 등을 설정해 멀티모달 API 호출을 구성했다.')
doc.add_paragraph('· 이미지·텍스트 혼합 프롬프트로 음식명·설명을 생성하고, 시간 제한/에러 처리를 적용했다.')

doc.add_heading('5. Food Matcher(FastAPI)', level=1)
doc.add_paragraph('· 모델: BM-K/KoSimCSE-roberta-multitask, SentenceTransformers')
doc.add_paragraph('· python embed_foods.py 로 foods 테이블에서 텍스트를 가져와 food_embeddings 테이블을 생성한다.')
doc.add_paragraph('· uvicorn server:app --port 9700 으로 API를 기동, /match 엔드포인트에서 상위 K개의 음식 코드와 유사도를 반환한다.')
doc.add_paragraph('· Laravel .env 의 FOOD_MATCHER_URL 이 이 서비스 주소를 가리키며, 실패 시 예외 처리를 수행한다.')

doc.add_heading('6. 데이터베이스 정제 및 전처리', level=1)
doc.add_paragraph('· 대표 식품명(common_name) 그룹 평균으로 지방(fat_g)과 탄수화물(carbohydrate_g)의 누락값을 채워 현실적인 수치로 보완했다.')
doc.add_paragraph('· product_weight 문자열을 파싱해 실제 제품 중량 기준으로 모든 영양 성분을 스케일링했다(예: 591g 포장 → 값 ×5.91).')
doc.add_paragraph('· 정제 후에도 common_name이 없는 일부 데이터는 후속 조사 대상임을 문서화했다.')

doc.add_heading('7. 운영 · 테스트 계획', level=1)
doc.add_paragraph('· 개발/테스트 환경: Node.js 18, PHP 8.2, MySQL(XAMPP), Python 3.11')
doc.add_paragraph('· 엔드투엔드 테스트: 로그인 → 사진 업로드 → 응답 JSON 확인(cURL/Postman) → React UI 렌더링 확인')
doc.add_paragraph('· FastAPI/임베딩 변경 시 embed_foods 재실행, Laravel 캐시 초기화 절차를 운영 문서에 포함했다.')

doc.add_heading('8. 향후 과제', level=1)
doc.add_paragraph('· eGL(estimated Glycemic Load) 모델 연동: 공개 데이터 조사 또는 자체 모델 학습으로 혈당 예측 기능 추가')
doc.add_paragraph('· 사용자 맞춤 추천: 섭취 이력 저장, 목표 칼로리 대비 가이던스 제공')
doc.add_paragraph('· 다국어 UI, 접근성 개선, 오프라인 캐시 기능')


doc.save(r'C:\AI\de_new.docx')
