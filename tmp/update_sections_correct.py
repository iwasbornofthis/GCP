from docx import Document

def remove_section(doc, start_prefix, stop_prefix):
    while True:
        paras = doc.paragraphs
        start_idx = None
        for i,p in enumerate(paras):
            if p.text.strip().startswith(start_prefix):
                start_idx = i
                break
        if start_idx is None:
            return None
        text = doc.paragraphs[start_idx].text.strip()
        if text.startswith(stop_prefix):
            return start_idx
        doc.paragraphs[start_idx]._element.getparent().remove(doc.paragraphs[start_idx]._element)
        

def insert_content(anchor, blocks):
    for style, text in blocks:
        para = anchor.insert_paragraph_before(text)
        if style:
            para.style = style


doc = Document(r"C:\AI\de.docx")

start4 = remove_section(doc, "4.", "5.")
anchor4 = doc.paragraphs[start4]
content4 = [
    ('Heading 1', '4. 핵심 기능 구성'),
    ('Heading 2', '4.1 이미지 업로드 및 전처리'),
    ('Normal', 'React 프런트엔드는 카메라 촬영과 갤러리 선택을 지원하고, 업로드 전에 파일 용량·확장자·EXIF 정보를 검증합니다. 모바일 브라우저에서 FormData를 구성해 Laravel API로 전송하며, 오류 발생 시 즉시 사용자 안내를 제공합니다.'),
    ('Heading 2', '4.2 OpenAI + Food Matcher 연동'),
    ('Normal', 'Laravel는 수신한 이미지를 저장한 뒤 OpenAI GPT-4o 멀티모달 API에 전달해 음식명·설명을 얻고, 동일한 텍스트를 Food Matcher FastAPI에 보내 임베딩 유사도로 food_code와 score를 계산합니다.'),
    ('Heading 2', '4.3 영양 정보 응답 구성'),
    ('Normal', '매칭된 food_code로 FOODS 테이블을 조회해 열량, 탄수화물, 단백질, 지방, 당, 나트륨 등을 가져오며, product_weight 환산과 common_name 평균 보정이 적용된 최신 값을 JSON으로 패키징합니다.'),
    ('Heading 2', '4.4 상태 관리 및 오류 대응'),
    ('Normal', 'OpenAI 타임아웃, 임베딩 미매칭, DB 누락 등 주요 예외에 대해 표준 응답 코드와 메시지를 정의했고, Laravel 로그·알림 훅으로 모니터링합니다. 사용자는 실패 사유와 재시도 방법을 즉시 확인할 수 있습니다.')
]
insert_content(anchor4, content4)

start5 = remove_section(doc, "5.", "6.")
anchor5 = doc.paragraphs[start5]
content5 = [
    ('Heading 1', '5. 사용자 여정'),
    ('Heading 2', '5.1 회원가입과 온보딩'),
    ('Normal', '① QR/링크로 접속하여 이메일·비밀번호로 회원가입한다. ② 나이·성별·키·체중·목표 칼로리를 입력하고 카메라 권한을 허용한다. ③ 튜토리얼에서 사진 업로드 흐름을 체험하며 첫 분석을 완료한다.'),
    ('Heading 2', '5.2 사진 분석 플로우'),
    ('Normal', '① 사용자가 사진을 촬영하거나 선택하면 React가 업로드 진행률을 표시한다. ② Laravel은 이미지를 검증 후 OpenAI와 Food Matcher를 순차 호출한다. ③ JSON 응답을 받아 영양 카드, 대체 식단, 오류 메시지를 렌더링한다.'),
    ('Heading 2', '5.3 결과 활용 및 기록'),
    ('Normal', '분석 결과는 MEAL_RECORDS에 저장되며 사용자는 일/주별 요약을 확인하거나 메모를 추가할 수 있다. 주간 리포트는 이메일/알림으로 전달되고, 누적 데이터는 eGL 모델 도입을 위한 학습 샘플로 축적된다.')
]
insert_content(anchor5, content5)


doc.save(r"C:\AI\de.docx")
